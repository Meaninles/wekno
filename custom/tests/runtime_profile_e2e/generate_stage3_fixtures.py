"""Generate a mixed-format, deterministic local ingestion batch."""

from __future__ import annotations

import argparse
import csv
import json
from pathlib import Path

from docx import Document
from openpyxl import Workbook
from reportlab.pdfgen import canvas


def body(marker: str, file_type: str, index: int) -> str:
    return (
        f"{marker} mixed-format concurrency probe. "
        f"This is {file_type} fixture {index}. "
        "Core parsing, summary generation, question generation, graph extraction, "
        "durable finalization, tenant fairness, and model-pool admission must remain isolated."
    )


def write_pdf(path: Path, text: str) -> None:
    page = canvas.Canvas(str(path))
    page.drawString(72, 760, text[:100])
    page.drawString(72, 740, text[100:200])
    page.save()


def write_docx(path: Path, text: str) -> None:
    document = Document()
    document.add_heading("Runtime profile stage 3", level=1)
    document.add_paragraph(text)
    document.save(path)


def write_xlsx(path: Path, text: str) -> None:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Acceptance"
    sheet.append(["marker", "content"])
    sheet.append([text.split()[0], text])
    workbook.save(path)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--output", required=True)
    parser.add_argument("--marker", required=True)
    parser.add_argument("--copies", type=int, default=3)
    args = parser.parse_args()

    output = Path(args.output).resolve()
    output.mkdir(parents=True, exist_ok=True)
    paths: list[str] = []

    for index in range(1, max(1, args.copies) + 1):
        fixtures = {
            "txt": output / f"stage3-{index:02d}.txt",
            "md": output / f"stage3-{index:02d}.md",
            "csv": output / f"stage3-{index:02d}.csv",
            "pdf": output / f"stage3-{index:02d}.pdf",
            "docx": output / f"stage3-{index:02d}.docx",
            "xlsx": output / f"stage3-{index:02d}.xlsx",
        }
        for file_type, path in fixtures.items():
            text = body(args.marker, file_type, index)
            if file_type == "txt":
                path.write_text(text, encoding="utf-8")
            elif file_type == "md":
                path.write_text(f"# Runtime profile stage 3\n\n{text}\n", encoding="utf-8")
            elif file_type == "csv":
                with path.open("w", encoding="utf-8", newline="") as stream:
                    writer = csv.writer(stream)
                    writer.writerow(["marker", "content"])
                    writer.writerow([args.marker, text])
            elif file_type == "pdf":
                write_pdf(path, text)
            elif file_type == "docx":
                write_docx(path, text)
            elif file_type == "xlsx":
                write_xlsx(path, text)
            paths.append(str(path))

    print(json.dumps(paths, ensure_ascii=False))


if __name__ == "__main__":
    main()
